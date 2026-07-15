"""
Document loaders.

Each loader takes a file path (+ optional source URL) and returns a list of
"documents", where every document is a dict:
    {"text": str, "metadata": {"source": str, "title": str, "page": int|None}}

Supported formats:
    * HTML  (.html)
    * PDF   (.pdf)  — with pdfplumber table extraction fallback
    * DOCX  (.docx)
    * TXT   (.txt)
"""
from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Optional

from bs4 import BeautifulSoup

from app.config import settings

# ------------------------------------------------------------------
# Noise filtering patterns
# ------------------------------------------------------------------
# PDF pages that are essentially blank or only contain page numbers / headers.
_NOISE_PATTERNS = re.compile(
    r"^("
    r"this page intentionally left blank\.?"
    r"|page \d+ of \d+"
    r"|\d+ / \d+"
    r"|\d+\s*$"
    r"|\*"
    r"|\.{3,}"
    r")$",
    re.IGNORECASE,
)

# Minimum meaningful text length — documents shorter than this are skipped.
MIN_DOC_LENGTH = 80


# ------------------------------------------------------------------
# Global nav deduplication
# ------------------------------------------------------------------
# Tracks nav content hashes so we index the navigation menu only ONCE
# across the entire manifest, not once per HTML page.
_seen_nav_hashes: set[str] = set()


def _reset_nav_dedup() -> None:
    """Reset the navigation dedup set (called at the start of each ingestion)."""
    global _seen_nav_hashes
    _seen_nav_hashes = set()


def _is_noise(text: str) -> bool:
    """Return True if the text is a noise page (blank, page number, etc.)."""
    stripped = text.strip()
    if len(stripped) < MIN_DOC_LENGTH:
        return True
    if _NOISE_PATTERNS.match(stripped):
        return True
    return False


# ------------------------------------------------------------------
# HTML
# ------------------------------------------------------------------
def _clean_text(text: str) -> str:
    """Collapse whitespace and blank lines."""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    return "\n".join(lines)


def _extract_tables(soup: BeautifulSoup) -> list[str]:
    """Extract HTML <table> elements as structured text.

    Many NIT Sikkim pages embed fee structures, schedules, and statistics
    in HTML tables.  Converting them to readable text ensures the data
    is indexed even when it sits outside the main content selectors.
    """
    table_texts: list[str] = []
    for table in soup.find_all("table"):
        rows: list[str] = []
        for tr in table.find_all("tr"):
            cells = [td.get_text(strip=True) for td in tr.find_all(["th", "td"])]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            table_texts.append("\n".join(rows))
    return table_texts


def load_html(path: Path, source_url: Optional[str] = None) -> list[dict]:
    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "lxml")

    title = (soup.title.string.strip() if soup.title and soup.title.string else path.stem)

    # Remove non-content noise.
    for tag in soup(["script", "style", "noscript", "form", "iframe"]):
        tag.decompose()

    docs: list[dict] = []

    # ---- 1. Extract NAVIGATION menu (deduplicated globally) ----
    # The NIT Sikkim site lists departments, facilities, etc. in collapsible
    # menu divs (mdBody). We index this ONCE globally so department names
    # are searchable without polluting the index with duplicates.
    nav_texts: list[str] = []
    for nav_div in soup.select("div.mdBody"):
        nav_texts.append(nav_div.get_text(separator="\n", strip=True))
    if nav_texts:
        nav_clean = _clean_text("\n".join(nav_texts))
        if len(nav_clean) > 50:
            nav_hash = hashlib.sha256(nav_clean.encode()).hexdigest()
            if nav_hash not in _seen_nav_hashes:
                _seen_nav_hashes.add(nav_hash)
                docs.append({
                    "text": f"Navigation menu / site structure of NIT Sikkim:\n{nav_clean}",
                    "metadata": {
                        "source": source_url or str(path),
                        "title": "NIT Sikkim Navigation Menu",
                        "type": "html",
                        "page": None,
                    },
                })

    # ---- 2. Extract the MAIN content area ----
    # Strategy (ordered by specificity):
    #   1. div.tab-contents  (NIT Sikkim's actual content container)
    #   2. <main> / <article> (semantic HTML5)
    #   3. Common CMS content divs
    #   4. <body> as last resort (with nav/footer/menus stripped)
    content_parts: list[str] = []

    # 1. NIT Sikkim specific: tab-contents holds the real page content.
    tab_contents = soup.select("div.tab-contents")
    if tab_contents:
        for tc in tab_contents:
            content_parts.append(tc.get_text(separator="\n", strip=True))

    # 2. Semantic HTML5 elements.
    if not content_parts:
        for tag_name in ["main", "article"]:
            found = soup.find(tag_name)
            if found:
                content_parts.append(found.get_text(separator="\n", strip=True))
                break

    # 3. Common CMS content containers.
    if not content_parts:
        for selector in ["#content", ".content", "#main-content", ".main-content",
                         ".page-content", "#page-content", ".body-content",
                         ".post-content", ".entry-content"]:
            found = soup.select_one(selector)
            if found:
                content_parts.append(found.get_text(separator="\n", strip=True))
                break

    # 4. Fallback: strip nav/header/footer/menus, then extract from the
    #    first content heading onward.  The NIT Sikkim nav lives in plain
    #    <div> overlays (not <nav> tags), so stripping standard elements
    #    alone isn't enough — but the nav never uses <h1>/<h2>/<h3>, so
    #    jumping to the first real heading cleanly skips all chrome.
    if not content_parts:
        body = soup.body or soup
        for tag in body.find_all(["nav", "header", "footer", "aside"]):
            tag.decompose()
        for tag in body.select("div.mdBody, #footer, #header, #nav, .navbar, .navigation"):
            tag.decompose()
        first_heading = body.find(["h1", "h2", "h3"])
        if first_heading:
            # Walk up to the top-level container (direct child of body)
            # that holds the heading, then take it + all following siblings.
            container = first_heading
            while container.parent is not None and container.parent.name not in ("body", "[document]"):
                container = container.parent
            parts: list[str] = [container.get_text(separator="\n", strip=True)]
            for sib in container.find_next_siblings():
                t = sib.get_text(separator="\n", strip=True)
                if t:
                    parts.append(t)
            content_parts.append("\n".join(p for p in parts if p))
        else:
            content_parts.append(body.get_text(separator="\n", strip=True))

    # ---- 3. Extract embedded tables ----
    # Tables often contain the actual data (fees, schedules, statistics)
    # that the text extraction above might flatten or miss.
    table_texts = _extract_tables(soup)
    if table_texts:
        content_parts.extend(table_texts)

    raw = "\n\n".join(p for p in content_parts if p)
    clean = _clean_text(raw)

    if clean and len(clean) >= 20:
        docs.append({
            "text": clean,
            "metadata": {
                "source": source_url or str(path),
                "title": title,
                "type": "html",
                "page": None,
            },
        })

    return docs


# ------------------------------------------------------------------
# PDF — with pdfplumber table extraction fallback
# ------------------------------------------------------------------
def _try_pdfplumber_tables(path: Path) -> list[str]:
    """Extract tables from a PDF using pdfplumber (if installed).

    Returns a list of table strings. Each table is formatted as rows
    of pipe-separated cells.
    """
    try:
        import pdfplumber
    except ImportError:
        return []

    table_texts: list[str] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                if not tables:
                    continue
                for table in tables:
                    rows: list[str] = []
                    for row in table:
                        cells = [str(c).strip() if c else "" for c in row]
                        if any(cells):
                            rows.append(" | ".join(cells))
                    if rows:
                        table_texts.append("\n".join(rows))
    except Exception:  # noqa: BLE001
        pass
    return table_texts


def load_pdf(path: Path, source_url: Optional[str] = None) -> list[dict]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    docs: list[dict] = []
    title = path.stem

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = text.strip()
        if not text:
            continue
        # Skip noise pages (blank, page-number-only, etc.)
        if _is_noise(text):
            continue
        docs.append({
            "text": text,
            "metadata": {
                "source": source_url or str(path),
                "title": title,
                "type": "pdf",
                "page": i + 1,
            },
        })

    # Also try pdfplumber for tables that pypdf might miss or flatten.
    table_texts = _try_pdfplumber_tables(path)
    if table_texts:
        for j, table_text in enumerate(table_texts):
            if _is_noise(table_text):
                continue
            docs.append({
                "text": f"[Table from {title}]\n{table_text}",
                "metadata": {
                    "source": source_url or str(path),
                    "title": f"{title} (Table {j + 1})",
                    "type": "pdf",
                    "page": None,
                },
            })

    return docs


# ------------------------------------------------------------------
# DOCX
# ------------------------------------------------------------------
def load_docx(path: Path, source_url: Optional[str] = None) -> list[dict]:
    import docx  # python-docx

    doc = docx.Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)

    docs: list[dict] = []

    if text and not _is_noise(text):
        docs.append({
            "text": text,
            "metadata": {
                "source": source_url or str(path),
                "title": path.stem,
                "type": "docx",
                "page": None,
            },
        })

    # Also extract tables from DOCX
    for i, table in enumerate(doc.tables):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        table_text = "\n".join(rows)
        if table_text and not _is_noise(table_text):
            docs.append({
                "text": f"[Table from {path.stem}]\n{table_text}",
                "metadata": {
                    "source": source_url or str(path),
                    "title": f"{path.stem} (Table {i + 1})",
                    "type": "docx",
                    "page": None,
                },
            })

    return docs


# ------------------------------------------------------------------
# TXT
# ------------------------------------------------------------------
def load_txt(path: Path, source_url: Optional[str] = None) -> list[dict]:
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    if not text or _is_noise(text):
        return []
    return [{
        "text": text,
        "metadata": {
            "source": source_url or str(path),
            "title": path.stem,
            "type": "txt",
            "page": None,
        },
    }]


# ------------------------------------------------------------------
# Dispatcher
# ------------------------------------------------------------------
_LOADERS = {
    ".html": load_html,
    ".htm": load_html,
    ".php": load_html,   # NIT Sikkim serves pages as .php — they're HTML
    ".asp": load_html,
    ".aspx": load_html,
    ".jsp": load_html,
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".doc": load_docx,   # best-effort; may fail on old .doc
    ".txt": load_txt,
}


def load_file(path: Path, source_url: Optional[str] = None) -> list[dict]:
    """Load any supported file by extension. Returns [] for unsupported types."""
    ext = path.suffix.lower()
    loader = _LOADERS.get(ext)
    if loader is None:
        return []
    try:
        return loader(path, source_url)
    except Exception as exc:  # noqa: BLE001
        print(f"[WARN] Failed to load {path}: {exc}")
        return []


def _enrich_metadata(docs: list[dict]) -> None:
    """Add is_archived, last_updated, and source_type to metadata."""
    # Look for years like 2010 to 2039 in title or source
    year_pattern = re.compile(r"(20[1-3][0-9])")
    
    for doc in docs:
        meta = doc["metadata"]
        
        meta["is_archived"] = False
        meta["last_updated"] = "2026-01-01"  # Default to current year
        meta["source_type"] = "webpage" if meta.get("type") in ("html", "php") else meta.get("type", "unknown")
        
        # Determine recency from title or source URL for documents
        if meta.get("type") in ("pdf", "docx", "doc"):
            title = meta.get("title", "")
            source = meta.get("source", "")
            
            years = [int(y) for y in year_pattern.findall(title + " " + source)]
            if years:
                max_year = max(years)
                meta["last_updated"] = f"{max_year}-01-01"
                if max_year < 2025:
                    meta["is_archived"] = True


def load_all_from_manifest() -> list[dict]:
    """
    Read the crawler manifest and load every referenced file.
    Returns a flat list of documents ready for chunking.

    Navigation menus are deduplicated globally — only the first unique
    navigation block is kept, not one per HTML page.
    """
    import json

    # Reset nav dedup for a fresh ingestion run.
    _reset_nav_dedup()

    manifest_path = settings.raw_path / "manifest.json"
    if not manifest_path.exists():
        raise FileNotFoundError(
            f"Manifest not found at {manifest_path}. Run the crawler first: python -m app.crawler"
        )

    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    all_docs: list[dict] = []

    for entry in manifest.get("html_pages", []) + manifest.get("documents", []):
        file_path = Path(entry["file"])
        url = entry.get("url")
        if not file_path.exists():
            continue
        docs = load_file(file_path, source_url=url)
        all_docs.extend(docs)

    _enrich_metadata(all_docs)

    return all_docs
